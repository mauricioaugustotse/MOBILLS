# -*- coding: utf-8 -*-

"""
Lê uma planilha (Excel ou CSV), processa despesas e gera um relatório Word.

Regras desta versão:
- Excel: usa sempre a 2ª aba (índice 1); falha se o arquivo tiver menos de 2 abas.
- Remove linhas não-transação (vazias, totais e valor inválido).
- Mantém regras de cálculo de Parcela/% do fluxo existente.
- Aplica cores na coluna Parcela: vermelho para negativo, azul para positivo (tags carol/m&c).
- Inclui gráfico donut de 2 níveis (categoria + subdivisão por tag).
"""

import os
import re
import tempfile
import zipfile
from xml.etree import ElementTree as ET

import docx
import numpy as np
import pandas as pd
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

COLUNAS_ESPERADAS = ["Data", "Descrição", "Conta", "Categoria", "Tags", "Valor", "%", "Parcela", "Situação"]
TAGS_GRAFICO = ["mauricio", "carol", "m&c"]
TAGS_TOTAL_PARCELA = ["carol", "m&c"]
TEMPLATE_PADRAO = "templates/FEV_2026_base.docx"
PARTES_LAYOUT_RIGIDAS = [
    "word/styles.xml",
    "word/settings.xml",
    "word/theme/theme1.xml",
    "word/numbering.xml",
]

WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_NS = {"w": WML_NS}
W_VAL = f"{{{WML_NS}}}val"


# --- Helpers para forçar larguras fixas no Word (tblGrid + tcW) ---
def _twips_from_cm(cm_value: float) -> int:
    # 1 cm ≈ 566.929 twips
    return int(round(float(cm_value) * 567))


def _ensure_tblPr(table):
    """Garante a existência de w:tblPr e o retorna."""
    tbl = table._element
    tbl_pr_list = tbl.xpath("./w:tblPr")
    if tbl_pr_list:
        return tbl_pr_list[0]
    tbl_pr = OxmlElement("w:tblPr")
    tbl.insert(0, tbl_pr)
    return tbl_pr


def _set_table_layout_fixed(table):
    """Força w:tblLayout type='fixed' dentro de w:tblPr."""
    tbl_pr = _ensure_tblPr(table)
    for child in list(tbl_pr):
        if child.tag == qn("w:tblLayout"):
            tbl_pr.remove(child)
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)


def _set_cell_width(cell, cm_value: float):
    """Define a largura da célula via w:tcW (dxa) e mantém cell.width (Cm)."""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for child in list(tcPr):
            if child.tag == qn("w:tcW"):
                tcPr.remove(child)
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:type"), "dxa")
        tcW.set(qn("w:w"), str(_twips_from_cm(cm_value)))
        tcPr.append(tcW)
        cell.width = Cm(cm_value)
    except Exception:
        cell.width = Cm(cm_value)


def _apply_table_grid(table, column_names, widths_map_cm: dict):
    """Define w:tblGrid com larguras em twips e ajusta w:tblW."""
    tbl = table._element
    tbl_pr = _ensure_tblPr(table)

    for grid in tbl.xpath("./w:tblGrid"):
        tbl.remove(grid)

    grid = OxmlElement("w:tblGrid")
    total_twips = 0
    for name in column_names:
        width_cm = float(widths_map_cm.get(name, 2.0))
        twips = _twips_from_cm(width_cm)
        total_twips += twips
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(twips))
        grid.append(grid_col)

    try:
        tbl.insert(1, grid)
    except Exception:
        tbl.append(grid)

    for child in list(tbl_pr):
        if child.tag == qn("w:tblW"):
            tbl_pr.remove(child)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_pr.append(tbl_w)


def _formatar_numero_br(valor):
    return f"{float(valor):,.2f}".replace(",", "TEMP").replace(".", ",").replace("TEMP", ".")


def _is_number(valor):
    return isinstance(valor, (int, float, np.integer, np.floating)) and not pd.isna(valor)


def _limpar_texto(valor):
    if pd.isna(valor):
        return ""
    return str(valor).strip()


def _normalizar_tag(valor):
    return _limpar_texto(valor).lower()


def set_cell_font(cell, text, bold=False, size=14, color=None, align="left"):
    p = cell.paragraphs[0]
    p.text = text
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.runs[0]
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _normalizar_caminho_abs(caminho):
    return os.path.normcase(os.path.abspath(os.path.realpath(caminho)))


def _resolver_caminho_template(caminho_template=None):
    if caminho_template:
        return _normalizar_caminho_abs(caminho_template)

    base_script = os.path.dirname(os.path.abspath(__file__))
    return _normalizar_caminho_abs(os.path.join(base_script, TEMPLATE_PADRAO))


def _limpar_linhas_tabela_mantendo_cabecalho(table):
    if len(table.rows) <= 1:
        return

    for row_idx in range(len(table.rows) - 1, 0, -1):
        row = table.rows[row_idx]
        row._tr.getparent().remove(row._tr)


def _remover_paragrafo(paragraph):
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def _limpar_blocos_grafico_existentes(doc):
    texto_subtitulo = "Distribuição de despesas por categoria e pessoa"
    for paragraph in list(doc.paragraphs)[::-1]:
        tem_drawing = bool(paragraph._element.xpath(".//w:drawing"))
        eh_subtitulo = paragraph.text.strip() == texto_subtitulo
        if tem_drawing or eh_subtitulo:
            _remover_paragrafo(paragraph)


def _atributos_ordenados(elem):
    if elem is None:
        return None
    return tuple(sorted((chave, valor) for chave, valor in elem.attrib.items()))


def _coletar_invariantes_documento(xml_bytes):
    root = ET.fromstring(xml_bytes)
    body = root.find("w:body", W_NS)
    if body is None:
        raise RuntimeError("document.xml inválido: w:body ausente.")

    sect_pr = body.find("w:sectPr", W_NS)
    if sect_pr is None:
        sect_pr = body.find(".//w:sectPr", W_NS)
    if sect_pr is None:
        raise RuntimeError("document.xml inválido: w:sectPr ausente.")

    pg_sz = sect_pr.find("w:pgSz", W_NS)
    pg_mar = sect_pr.find("w:pgMar", W_NS)

    primeiro_paragrafo = body.find("w:p", W_NS)
    if primeiro_paragrafo is None:
        raise RuntimeError("document.xml inválido: primeiro parágrafo ausente.")
    estilo_titulo = primeiro_paragrafo.find("w:pPr/w:pStyle", W_NS)
    estilo_titulo_valor = estilo_titulo.attrib.get(W_VAL) if estilo_titulo is not None else None

    primeira_tabela = body.find("w:tbl", W_NS)
    if primeira_tabela is None:
        raise RuntimeError("document.xml inválido: primeira tabela ausente.")
    tbl_pr = primeira_tabela.find("w:tblPr", W_NS)
    if tbl_pr is None:
        raise RuntimeError("document.xml inválido: w:tblPr ausente na primeira tabela.")

    tbl_style = tbl_pr.find("w:tblStyle", W_NS)
    tbl_layout = tbl_pr.find("w:tblLayout", W_NS)
    tbl_w = tbl_pr.find("w:tblW", W_NS)
    tbl_ind = tbl_pr.find("w:tblInd", W_NS)
    tbl_grid = primeira_tabela.find("w:tblGrid", W_NS)

    grid_cols = []
    if tbl_grid is not None:
        grid_cols = [_atributos_ordenados(col) for col in tbl_grid.findall("w:gridCol", W_NS)]

    return {
        "section_pgSz": _atributos_ordenados(pg_sz),
        "section_pgMar": _atributos_ordenados(pg_mar),
        "titulo_pStyle": estilo_titulo_valor,
        "tblStyle": _atributos_ordenados(tbl_style),
        "tblLayout": _atributos_ordenados(tbl_layout),
        "tblW": _atributos_ordenados(tbl_w),
        "tblInd": _atributos_ordenados(tbl_ind),
        "tblGridCols": tuple(grid_cols),
    }


def _normalizar_xml_canonico(xml_bytes):
    try:
        return ET.canonicalize(xml_bytes.decode("utf-8")).encode("utf-8")
    except Exception:
        try:
            root = ET.fromstring(xml_bytes)
            return ET.tostring(root, encoding="utf-8")
        except Exception:
            return xml_bytes


def _validar_layout_rigido(caminho_template, caminho_gerado):
    for parte in PARTES_LAYOUT_RIGIDAS:
        try:
            with zipfile.ZipFile(caminho_template, "r") as zip_template:
                conteudo_template = zip_template.read(parte)
        except KeyError as exc:
            raise RuntimeError(f"Template inválido: parte ausente '{parte}'.") from exc

        try:
            with zipfile.ZipFile(caminho_gerado, "r") as zip_saida:
                conteudo_saida = zip_saida.read(parte)
        except KeyError as exc:
            raise RuntimeError(f"Saída inválida: parte ausente '{parte}'.") from exc

        if _normalizar_xml_canonico(conteudo_template) != _normalizar_xml_canonico(conteudo_saida):
            raise RuntimeError(f"Validação rígida falhou: parte diferente '{parte}'.")

    with zipfile.ZipFile(caminho_template, "r") as zip_template:
        doc_template_xml = zip_template.read("word/document.xml")
    with zipfile.ZipFile(caminho_gerado, "r") as zip_saida:
        doc_saida_xml = zip_saida.read("word/document.xml")

    inv_template = _coletar_invariantes_documento(doc_template_xml)
    inv_saida = _coletar_invariantes_documento(doc_saida_xml)

    divergencias = []
    for chave in [
        "section_pgSz",
        "section_pgMar",
        "titulo_pStyle",
        "tblStyle",
        "tblLayout",
        "tblW",
        "tblInd",
        "tblGridCols",
    ]:
        if inv_template.get(chave) != inv_saida.get(chave):
            divergencias.append(chave)

    if divergencias:
        lista = ", ".join(divergencias)
        raise RuntimeError(
            "Validação rígida falhou em invariantes de document.xml: "
            f"{lista}."
        )


def _verificar_dependencia_excel(caminho_arquivo):
    extensao = os.path.splitext(caminho_arquivo)[1].lower()
    if extensao in (".xlsx", ".xlsm", ".xltx", ".xltm"):
        import importlib.util as _imp_util
        if _imp_util.find_spec("openpyxl") is None:
            return "openpyxl"
    elif extensao == ".xls":
        import importlib.util as _imp_util
        if _imp_util.find_spec("xlrd") is None:
            return "xlrd"
    return None


def _verificar_dependencia_grafico():
    import importlib.util as _imp_util
    if _imp_util.find_spec("matplotlib") is None:
        return "matplotlib"
    return None


def _obter_titulo_mes(nome_base):
    meses_pt = {
        "JAN": "Janeiro",
        "FEV": "Fevereiro",
        "MAR": "Março",
        "ABR": "Abril",
        "MAI": "Maio",
        "JUN": "Junho",
        "JUL": "Julho",
        "AGO": "Agosto",
        "SET": "Setembro",
        "OUT": "Outubro",
        "NOV": "Novembro",
        "DEZ": "Dezembro",
    }
    match = re.match(r"^([A-Za-z]{3})_(\d{4})$", nome_base.strip())
    if not match:
        return nome_base

    mes_abrev = match.group(1).upper()
    ano = match.group(2)
    if mes_abrev in meses_pt:
        return f"{meses_pt[mes_abrev]} {ano}"
    return nome_base


def _ler_dataframe_excel_segunda_aba(caminho_arquivo):
    dependencia_faltando = _verificar_dependencia_excel(caminho_arquivo)
    if dependencia_faltando:
        raise RuntimeError(
            "ERRO FATAL ao tentar ler a planilha: "
            f"Dependência '{dependencia_faltando}' não instalada. "
            f"Instale com: pip install {dependencia_faltando}"
        )

    excel_file = pd.ExcelFile(caminho_arquivo)
    if len(excel_file.sheet_names) < 2:
        raise RuntimeError(
            "ERRO FATAL: arquivo Excel precisa ter pelo menos 2 abas. "
            "Este fluxo usa sempre a 2ª aba (índice 1)."
        )

    aba_escolhida = excel_file.sheet_names[1]
    dataframe = pd.read_excel(caminho_arquivo, sheet_name=aba_escolhida)
    return dataframe, aba_escolhida, excel_file.sheet_names


def _filtrar_linhas_nao_transacao(df):
    colunas_base = ["Data", "Descrição", "Valor", "Conta", "Situação", "Categoria", "Tags"]
    for col in colunas_base:
        if col not in df.columns:
            df[col] = np.nan

    total_inicial = len(df)
    df_filtrado = df.dropna(how="all", subset=colunas_base).copy()
    removidas_vazias = total_inicial - len(df_filtrado)

    data_str = df_filtrado["Data"].astype(str).str.strip().str.lower()
    mask_total = data_str.str.startswith("total")
    removidas_total = int(mask_total.sum())
    df_filtrado = df_filtrado.loc[~mask_total].copy()

    df_filtrado["Valor"] = pd.to_numeric(df_filtrado["Valor"], errors="coerce")
    mask_valor_invalido = df_filtrado["Valor"].isna()
    removidas_valor_invalido = int(mask_valor_invalido.sum())
    df_filtrado = df_filtrado.loc[~mask_valor_invalido].copy()

    return df_filtrado, {
        "inicial": total_inicial,
        "removidas_vazias": removidas_vazias,
        "removidas_total": removidas_total,
        "removidas_valor_invalido": removidas_valor_invalido,
        "final": len(df_filtrado),
    }


def _extrair_total_linha_total(df_original):
    if df_original is None:
        return None

    if "Data" not in df_original.columns or "Valor" not in df_original.columns:
        return None

    data_str = df_original["Data"].astype(str).str.strip().str.lower()
    mask_total = data_str.str.startswith("total")
    if not mask_total.any():
        return None

    valores_total = pd.to_numeric(df_original.loc[mask_total, "Valor"], errors="coerce").dropna()
    if valores_total.empty:
        return None

    return float(valores_total.iloc[-1])


def _ajustar_luminosidade(rgb, fator):
    if fator < 1:
        return tuple(max(0.0, min(1.0, c * fator)) for c in rgb)

    ganho = fator - 1
    return tuple(max(0.0, min(1.0, c + (1 - c) * ganho)) for c in rgb)


def _gerar_grafico_donut_categoria_tag(df_processado, caminho_png, total_referencia=None):
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.colors import to_rgb
    from matplotlib.patches import Patch

    df_chart = df_processado.copy()
    df_chart["Tags"] = df_chart["Tags"].apply(_normalizar_tag)
    df_chart["Categoria"] = df_chart["Categoria"].apply(_limpar_texto)
    df_chart["Valor"] = pd.to_numeric(df_chart["Valor"], errors="coerce")

    df_chart = df_chart[df_chart["Tags"].isin(TAGS_GRAFICO)].copy()
    df_chart = df_chart[df_chart["Categoria"] != ""].copy()
    df_chart = df_chart[df_chart["Valor"].notna()].copy()
    df_chart["metric"] = df_chart["Valor"].abs()
    df_chart = df_chart[df_chart["metric"] > 0].copy()

    if df_chart.empty:
        raise RuntimeError("ERRO FATAL: não há dados válidos para gerar o gráfico donut.")

    limite_outros = 500.0
    totais_categoria_original = df_chart.groupby("Categoria", as_index=True)["metric"].sum()
    categorias_para_outros = totais_categoria_original[totais_categoria_original < limite_outros].index.tolist()
    if categorias_para_outros:
        df_chart["Categoria_Ajustada"] = df_chart["Categoria"].where(
            ~df_chart["Categoria"].isin(categorias_para_outros),
            "outros",
        )
    else:
        df_chart["Categoria_Ajustada"] = df_chart["Categoria"]

    totais_categoria = (
        df_chart.groupby("Categoria_Ajustada", as_index=True)["metric"].sum().sort_values(ascending=False)
    )
    if "outros" in totais_categoria.index and len(totais_categoria.index) > 1:
        valor_outros = float(totais_categoria.loc["outros"])
        totais_categoria = totais_categoria.drop(index="outros")
        totais_categoria.loc["outros"] = valor_outros

    subdivisoes = df_chart.groupby(["Categoria_Ajustada", "Tags"], as_index=True)["metric"].sum()

    categorias = totais_categoria.index.tolist()
    # Paleta de alto contraste para diferenciar melhor as categorias.
    paleta_contraste = [
        "#1f77b4",
        "#ff7f0e",
        "#2ca02c",
        "#d62728",
        "#9467bd",
        "#8c564b",
        "#e377c2",
        "#7f7f7f",
        "#bcbd22",
        "#17becf",
        "#393b79",
        "#637939",
        "#8c6d31",
        "#843c39",
        "#7b4173",
    ]
    mapa_cor_base = {}
    indice_paleta = 0
    for categoria in categorias:
        if categoria == "outros":
            mapa_cor_base[categoria] = to_rgb("#9e9e9e")
        else:
            mapa_cor_base[categoria] = to_rgb(paleta_contraste[indice_paleta % len(paleta_contraste)])
            indice_paleta += 1

    categoria_valores = totais_categoria.values.tolist()
    categoria_cores = [mapa_cor_base[categoria] for categoria in categorias]

    estilos_tag = {
        "m&c": {"fator": 0.72, "hatch": "///", "edge_factor": 0.75, "linewidth": 1.0},
        "carol": {"fator": 1.00, "hatch": "", "edge_factor": 0.75, "linewidth": 1.0},
        "mauricio": {"fator": 1.22, "hatch": "..", "edge_factor": 0.75, "linewidth": 1.0},
    }
    ordem_tags_tom = ["m&c", "carol", "mauricio"]
    subdiv_valores = []
    subdiv_cores = []
    subdiv_estilos = []
    for categoria in categorias:
        cor_base = mapa_cor_base[categoria]
        if categoria == "outros":
            valor_outros = float(totais_categoria.loc[categoria])
            if valor_outros > 0:
                subdiv_valores.append(valor_outros)
                subdiv_cores.append(cor_base)
                subdiv_estilos.append({"hatch": "", "edge_factor": 0.82, "linewidth": 1.0})
            continue

        for tag in ordem_tags_tom:
            valor = float(subdivisoes.get((categoria, tag), 0.0))
            if valor <= 0:
                continue
            estilo = estilos_tag[tag]
            subdiv_valores.append(valor)
            subdiv_cores.append(_ajustar_luminosidade(cor_base, estilo["fator"]))
            subdiv_estilos.append(estilo)

    if not subdiv_valores:
        raise RuntimeError("ERRO FATAL: não foi possível montar as subdivisões do gráfico donut.")

    fig, ax = plt.subplots(figsize=(14, 8), dpi=180)

    wedges_inner, _ = ax.pie(
        subdiv_valores,
        radius=1.0,
        colors=subdiv_cores,
        startangle=90,
        counterclock=False,
        wedgeprops={"width": 0.30, "edgecolor": "white", "linewidth": 0.8},
    )
    for wedge, cor_atual, estilo in zip(wedges_inner, subdiv_cores, subdiv_estilos):
        if estilo["hatch"]:
            wedge.set_hatch(estilo["hatch"])
        wedge.set_edgecolor(_ajustar_luminosidade(cor_atual, estilo["edge_factor"]))
        wedge.set_linewidth(estilo["linewidth"])

    ax.pie(
        categoria_valores,
        radius=1.34,
        colors=categoria_cores,
        startangle=90,
        counterclock=False,
        wedgeprops={"width": 0.34, "edgecolor": "white", "linewidth": 1.2},
    )

    if _is_number(total_referencia):
        total = abs(float(total_referencia))
    else:
        total = abs(float(pd.to_numeric(df_processado["Valor"], errors="coerce").dropna().sum()))

    ax.text(
        0,
        0,
        f"Total\nR$ {_formatar_numero_br(total)}",
        ha="center",
        va="center",
        fontsize=12,
        fontweight="bold",
    )

    legenda_categoria = [
        Patch(facecolor=categoria_cores[i], label=f"{cat} (R$ {_formatar_numero_br(categoria_valores[i])})")
        for i, cat in enumerate(categorias)
    ]
    cor_exemplo_tom = to_rgb("#4e79a7")
    legenda_tag = [
        Patch(
            facecolor=_ajustar_luminosidade(cor_exemplo_tom, estilos_tag["m&c"]["fator"]),
            edgecolor=_ajustar_luminosidade(cor_exemplo_tom, estilos_tag["m&c"]["edge_factor"]),
            hatch=estilos_tag["m&c"]["hatch"],
            label="m&c (tom escuro)",
        ),
        Patch(
            facecolor=_ajustar_luminosidade(cor_exemplo_tom, estilos_tag["carol"]["fator"]),
            edgecolor=_ajustar_luminosidade(cor_exemplo_tom, estilos_tag["carol"]["edge_factor"]),
            label="carol (tom base)",
        ),
        Patch(
            facecolor=_ajustar_luminosidade(cor_exemplo_tom, estilos_tag["mauricio"]["fator"]),
            edgecolor=_ajustar_luminosidade(cor_exemplo_tom, estilos_tag["mauricio"]["edge_factor"]),
            hatch=estilos_tag["mauricio"]["hatch"],
            label="mauricio (tom claro)",
        ),
    ]

    legenda_categoria_artist = ax.legend(
        handles=legenda_categoria,
        title="Categorias",
        loc="center left",
        bbox_to_anchor=(1.028, 0.50),
        fontsize=13,
        title_fontsize=14,
        frameon=False,
        borderaxespad=0.12,
    )
    ax.add_artist(legenda_categoria_artist)

    ax.legend(
        handles=legenda_tag,
        title="Subdivisão por pessoa",
        loc="upper center",
        bbox_to_anchor=(0.5, -0.05),
        ncol=3,
        fontsize=13,
        title_fontsize=14,
        frameon=False,
    )

    ax.set_title("Distribuição de despesas por categoria e pessoa", fontsize=13, pad=30)
    fig.subplots_adjust(left=0.03, right=0.60, top=0.82, bottom=0.20)
    fig.savefig(caminho_png, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def selecionar_arquivo():
    """Abre o seletor de arquivos para escolher a planilha de entrada."""
    try:
        from tkinter import Tk, filedialog

        root = Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        caminho = filedialog.askopenfilename(
            title="Selecione a planilha de entrada",
            filetypes=[
                ("Planilhas Excel", "*.xlsx *.xls"),
                ("Arquivos CSV", "*.csv"),
                ("Todos os arquivos", "*.*"),
            ],
        )
        root.destroy()
        return caminho if caminho else None
    except Exception as exc:
        print(f"ERRO ao abrir o seletor de arquivos: {exc}")
        return None


def processar_e_gerar_docx(caminho_arquivo, verbose=False, caminho_template=None, validacao_layout="rigida"):
    """
    Executa leitura, cálculo e geração do relatório .docx.

    Args:
        caminho_arquivo (str): caminho do arquivo de entrada (.xlsx, .xls ou .csv).
        verbose (bool): ativa logs detalhados.
        caminho_template (str | None): caminho do template .docx base de layout.
        validacao_layout (str): "rigida" para validar invariantes de layout ou "off" para desativar.
    """
    if not os.path.exists(caminho_arquivo):
        print(f"ERRO: O arquivo '{caminho_arquivo}' não foi encontrado.")
        return

    modo_validacao = _limpar_texto(validacao_layout).lower() or "rigida"
    if modo_validacao not in ("rigida", "off"):
        print("ERRO: parâmetro validacao_layout inválido. Use 'rigida' ou 'off'.")
        return

    caminho_template_efetivo = _resolver_caminho_template(caminho_template)
    if not os.path.exists(caminho_template_efetivo):
        print(f"ERRO FATAL: template não encontrado em '{caminho_template_efetivo}'.")
        return

    dependencia_grafico = _verificar_dependencia_grafico()
    if dependencia_grafico:
        print(
            "ERRO FATAL ao gerar relatório: "
            f"Dependência '{dependencia_grafico}' não instalada. "
            f"Instale com: pip install {dependencia_grafico}"
            
        )
        return

    nome_base = os.path.splitext(os.path.basename(caminho_arquivo))[0]
    titulo_relatorio = _obter_titulo_mes(nome_base)

    try:
        if caminho_arquivo.lower().endswith(".csv"):
            try:
                df = pd.read_csv(caminho_arquivo, sep=None, engine="python")
            except Exception:
                df = pd.read_csv(caminho_arquivo, sep=";")
            aba_origem = "CSV"
            if verbose:
                print("Entrada CSV: leitura direta sem seleção de abas.")
        else:
            df, aba_origem, abas_disponiveis = _ler_dataframe_excel_segunda_aba(caminho_arquivo)
            print(f"Aba utilizada (índice 1): '{aba_origem}'. Abas disponíveis: {abas_disponiveis}")

        total_referencia_planilha = _extrair_total_linha_total(df)
        if total_referencia_planilha is not None:
            print(f"Total de referência detectado na linha de total: {_formatar_numero_br(abs(total_referencia_planilha))}.")
        else:
            print("Linha de total não encontrada (ou inválida). Total do gráfico usará fallback pela soma de Valor.")

        for col in COLUNAS_ESPERADAS:
            if col not in df.columns:
                df[col] = np.nan

        df, info_filtro = _filtrar_linhas_nao_transacao(df)
        print(
            "Filtro de linhas concluído: "
            f"inicial={info_filtro['inicial']}, "
            f"vazias={info_filtro['removidas_vazias']}, "
            f"totais={info_filtro['removidas_total']}, "
            f"valor_inválido={info_filtro['removidas_valor_invalido']}, "
            f"final={info_filtro['final']}."
        )

    except Exception as exc:
        print(f"ERRO FATAL ao tentar ler/filtrar a planilha: {exc}")
        return

    if df.empty:
        print("ERRO FATAL: não há linhas válidas para processar após a filtragem.")
        return

    for col in ["Data", "Descrição", "Conta", "Categoria", "Tags", "Situação"]:
        df[col] = df[col].apply(_limpar_texto)

    df["Valor"] = pd.to_numeric(df["Valor"], errors="coerce")
    df = df.loc[df["Valor"].notna()].copy()
    if df.empty:
        print("ERRO FATAL: não há valores numéricos válidos na coluna 'Valor'.")
        return

    df["Parcela"] = df["Valor"].copy()
    df["%"] = "1"
    logs = []

    contas_excluidas_inversao = ["Itaú - C.Corrente", "Banco do Brasil - C.Corrente"]

    for index, row in df.iterrows():
        valor = row["Valor"]
        tags = _normalizar_tag(row["Tags"])
        conta = _limpar_texto(row["Conta"])
        situacao = _limpar_texto(row["Situação"])

        deve_inverter_sinal_negativo = False
        if situacao == "Paga":
            if conta not in contas_excluidas_inversao:
                deve_inverter_sinal_negativo = True

        if tags == "carol" and situacao == "Paga":
            novo_valor = abs(valor)
            logs.append(
                f"[linha {index:04d}] Regra: Paga, carol. "
                f"Valor {valor:8.2f} -> {novo_valor:8.2f} (invertido para POSITIVO)."
            )
            df.loc[index, "Parcela"] = novo_valor
        elif tags == "m&c":
            novo_valor = valor / 2
            df.loc[index, "%"] = "0,5"
            if deve_inverter_sinal_negativo:
                novo_valor = -novo_valor
                logs.append(
                    f"[linha {index:04d}] Regra: Paga, m&c. "
                    f"Valor {valor:8.2f} -> {novo_valor:8.2f} (invertido e dividido)."
                )
            else:
                logs.append(
                    f"[linha {index:04d}] Regra: m&c (sem inversão). "
                    f"Valor {valor:8.2f} -> {novo_valor:8.2f} (dividido)."
                )
            df.loc[index, "Parcela"] = novo_valor
        elif tags == "mauricio":
            novo_valor = valor
            if deve_inverter_sinal_negativo and valor < 0:
                novo_valor = -valor
                logs.append(
                    f"[linha {index:04d}] Regra: Paga, mauricio. "
                    f"Valor {valor:8.2f} -> {novo_valor:8.2f} (invertido para POSITIVO)."
                )
            else:
                logs.append(f"[linha {index:04d}] Regra: mauricio (sem inversão). Valor {valor:8.2f} mantido.")
            df.loc[index, "Parcela"] = novo_valor
        else:
            novo_valor = valor
            if deve_inverter_sinal_negativo and valor < 0:
                novo_valor = -valor
                logs.append(
                    f"[linha {index:04d}] Regra: Paga, outrem. "
                    f"Valor {valor:8.2f} -> {novo_valor:8.2f} (invertido para POSITIVO)."
                )
            else:
                logs.append(f"[linha {index:04d}] Regra: {tags} (sem inversão). Valor {valor:8.2f} mantido.")
            df.loc[index, "Parcela"] = novo_valor

    if verbose:
        for linha in logs[:20]:
            print(linha)
        if len(logs) > 20:
            print(f"... {len(logs) - 20} logs adicionais omitidos.")

    df["Conta"] = df["Conta"].str.replace("Itaú - C.Corrente", "Itaú", regex=False)
    df.sort_values(by=["Tags", "Valor"], ascending=[True, True], inplace=True)

    colunas_ordenadas = ["Data", "Descrição", "Conta", "Categoria", "Tags", "Valor", "%", "Parcela", "Situação"]
    df_final = df.reindex(columns=colunas_ordenadas).copy()

    soma_valor_original = float(df["Valor"].sum())
    filtro_parcela_tags = df["Tags"].apply(_normalizar_tag).isin(TAGS_TOTAL_PARCELA)
    soma_parcela_especifica = float(df.loc[filtro_parcela_tags, "Parcela"].sum())
    print(
        "Totais calculados: "
        f"soma Valor={_formatar_numero_br(soma_valor_original)}, "
        f"soma Parcela(carol+m&c)={_formatar_numero_br(soma_parcela_especifica)}."
    )

    try:
        doc = docx.Document(caminho_template_efetivo)
    except Exception as exc:
        print(f"ERRO FATAL ao abrir template '{caminho_template_efetivo}': {exc}")
        return

    if not doc.paragraphs:
        print("ERRO FATAL: template inválido, não há parágrafos para o título.")
        return

    doc.paragraphs[0].text = titulo_relatorio

    if not doc.tables:
        print("ERRO FATAL: template inválido, não há tabela principal.")
        return

    table = doc.tables[0]
    if len(table.rows) == 0:
        print("ERRO FATAL: template inválido, tabela principal sem cabeçalho.")
        return

    _limpar_linhas_tabela_mantendo_cabecalho(table)

    hdr_cells = table.rows[0].cells
    if len(hdr_cells) != len(df_final.columns):
        print(
            "ERRO FATAL: template incompatível. "
            f"Tabela possui {len(hdr_cells)} colunas, esperado {len(df_final.columns)}."
        )
        return

    for i, col_name in enumerate(df_final.columns):
        set_cell_font(hdr_cells[i], col_name, bold=True, align="center")

    for _, row in df_final.iterrows():
        row_cells = table.add_row().cells

        for i, col_name in enumerate(df_final.columns):
            font_color = None
            alinhamento = "justify"

            if col_name == "Parcela":
                tags = _normalizar_tag(row["Tags"])
                parcela = row["Parcela"]
                if tags in TAGS_TOTAL_PARCELA and _is_number(parcela):
                    if parcela < 0:
                        font_color = RGBColor(255, 0, 0)
                    elif parcela > 0:
                        font_color = RGBColor(0, 0, 255)

            valor_celula = row[col_name]
            if _is_number(valor_celula):
                texto_formatado = _formatar_numero_br(valor_celula)
                alinhamento = "right"
            else:
                texto_formatado = str(valor_celula if pd.notna(valor_celula) else "")
                if col_name == "Tags":
                    alinhamento = "center"

            if col_name in ("Conta", "%"):
                alinhamento = "center"

            set_cell_font(row_cells[i], texto_formatado, color=font_color, align=alinhamento)

    total_cells = table.add_row().cells

    valor_col_index = list(df_final.columns).index("Valor")
    parcela_col_index = list(df_final.columns).index("Parcela")

    total_parcela_color = None
    if soma_parcela_especifica < 0:
        total_parcela_color = RGBColor(238, 0, 0)
    elif soma_parcela_especifica > 0:
        total_parcela_color = RGBColor(0, 0, 255)

    set_cell_font(total_cells[0], "TOTAIS", bold=True, align="center")
    set_cell_font(total_cells[valor_col_index], _formatar_numero_br(soma_valor_original), bold=True, align="right")
    set_cell_font(
        total_cells[parcela_col_index],
        _formatar_numero_br(soma_parcela_especifica),
        bold=True,
        color=total_parcela_color,
        align="right",
    )

    _limpar_blocos_grafico_existentes(doc)
    doc.add_paragraph()
    subtitulo = doc.add_paragraph("Distribuição de despesas por categoria e pessoa - próxima página")
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitulo.runs:
        subtitulo.runs[0].font.size = Pt(20)

    caminho_grafico_tmp = None
    try:
        with tempfile.NamedTemporaryFile(prefix="grafico_donut_", suffix=".png", delete=False) as arquivo_tmp:
            caminho_grafico_tmp = arquivo_tmp.name

        _gerar_grafico_donut_categoria_tag(
            df,
            caminho_grafico_tmp,
            total_referencia=total_referencia_planilha,
        )
        doc.add_picture(caminho_grafico_tmp, width=Cm(20))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("Gráfico donut gerado e inserido no documento.")

    except Exception as exc:
        print(f"ERRO FATAL ao gerar/inserir gráfico: {exc}")
        if caminho_grafico_tmp and os.path.exists(caminho_grafico_tmp):
            os.remove(caminho_grafico_tmp)
        return

    nome_saida = os.path.splitext(os.path.basename(caminho_arquivo))[0]
    pasta_saida = os.path.dirname(caminho_arquivo)
    caminho_saida = os.path.join(pasta_saida, f"{nome_saida}.docx")
    caminho_tmp_docx = f"{caminho_saida}.tmp"
    if _normalizar_caminho_abs(caminho_saida) == _normalizar_caminho_abs(caminho_template_efetivo):
        print(
            "ERRO FATAL: o caminho de saída coincide com o template base. "
            "Altere o nome/local do arquivo de entrada ou informe outro template."
        )
        if caminho_grafico_tmp and os.path.exists(caminho_grafico_tmp):
            os.remove(caminho_grafico_tmp)
        return

    try:
        doc.save(caminho_tmp_docx)
        if modo_validacao == "rigida":
            _validar_layout_rigido(caminho_template_efetivo, caminho_tmp_docx)
            print("Validação rígida de layout concluída com sucesso.")
        try:
            if os.path.exists(caminho_saida):
                os.remove(caminho_saida)
        except Exception:
            pass
        os.replace(caminho_tmp_docx, caminho_saida)
        print(f"Arquivo Word gerado com sucesso: {caminho_saida}")

    except Exception as exc:
        print(f"ERRO FATAL ao salvar/validar o documento: {exc}")
        if os.path.exists(caminho_tmp_docx):
            os.remove(caminho_tmp_docx)
        return

    finally:
        if caminho_grafico_tmp and os.path.exists(caminho_grafico_tmp):
            os.remove(caminho_grafico_tmp)


if __name__ == "__main__":
    nome_do_arquivo_de_entrada = selecionar_arquivo()
    if nome_do_arquivo_de_entrada:
        print(f"\nArquivo selecionado: '{nome_do_arquivo_de_entrada}'. Iniciando processamento...")
        modo_verbose = True
        processar_e_gerar_docx(nome_do_arquivo_de_entrada, verbose=modo_verbose)
    else:
        print("Nenhum arquivo foi selecionado.")

    try:
        input("\nPressione Enter para sair...")
    except Exception:
        pass
